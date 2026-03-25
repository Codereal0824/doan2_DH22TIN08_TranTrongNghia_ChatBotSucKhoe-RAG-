"""
Document Loader - Đọc và xử lý tài liệu PDF, Word, Text
"""
from config.config import config
from pathlib import Path
from typing import List, Dict
import sys

# ============================================
# CƠ CHẾ NHẬP THƯ VIỆN DỰ PHÒNG (GRACEFUL DEGRADATION)
# ============================================
# Thay vì import trực tiếp gây lỗi hệ thống nếu môi trường thiếu thư viện,
# việc dùng try-except giúp hệ thống vẫn chạy được các file .txt cơ bản
# ngay cả khi chưa cài đặt pypdf hay python-docx.

# Xử lý PDF
try:
    from pypdf import PdfReader
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

# Xử lý Word (.docx)
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Thêm path để import config từ thư mục gốc dự án
sys.path.append(str(Path(__file__).parent.parent.parent))

# Lớp DocumentLoader đóng vai trò là một "Adapter" (Bộ chuyển đổi) chuẩn hóa dữ liệu.
# Nhiệm vụ của nó là đọc các định dạng tệp khác nhau (Unstructured Data)
# và thống nhất chúng về một định dạng cấu trúc duy nhất (Dictionary) để nạp vào RAG.


class DocumentLoader:
    """Class để load documents từ nhiều định dạng"""

    def __init__(self):
        """Khởi tạo loader"""
        self.supported_formats = []

        if HAS_PYPDF:
            self.supported_formats.append('.pdf')
        if HAS_DOCX:
            self.supported_formats.extend(['.docx', '.doc'])

        # Mặc định luôn hỗ trợ đọc văn bản thuần tủy (.txt, .md)
        self.supported_formats.extend(['.txt', '.md'])

        print("Document Loader san sang")
        print(f"Ho tro: {', '.join(self.supported_formats)}")

    def load_pdf(self, file_path: str) -> Dict:
        """
        Đọc file PDF

        Args:
            file_path: Đường dẫn file PDF

        Returns:
            Dict: {'content': str, 'metadata': dict}
        """
        if not HAS_PYPDF:
            raise ImportError("Chua cai dat pypdf! Run: pip install pypdf")

        try:
            reader = PdfReader(file_path)

            # Trích xuất toàn bộ văn bản (Text Extraction) từ các trang PDF.
            # Lưu ý: Pypdf chỉ trích xuất được text, không đọc được text nằm trong hình ảnh (cần OCR).
            content = ""
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text:
                    # Đánh dấu phân trang để dễ truy vết (Traceability) sau này
                    content += f"\n--- Trang {page_num} ---\n{text}"

            metadata = {
                'source': Path(file_path).name,
                'file_path': str(file_path),
                'type': 'pdf',
                'pages': len(reader.pages),
                'format': 'PDF'
            }

            # Trích xuất siêu dữ liệu (Metadata) ẩn bên trong file PDF nếu có
            if reader.metadata:
                if reader.metadata.title:
                    metadata['title'] = reader.metadata.title
                if reader.metadata.author:
                    metadata['author'] = reader.metadata.author

            print(
                f"Da doc PDF: {Path(file_path).name} ({len(reader.pages)} trang)")

            return {
                'content': content.strip(),
                'metadata': metadata
            }

        except Exception as e:
            print(f"Loi doc PDF {file_path}: {e}")
            return {'content': '', 'metadata': {'error': str(e)}}

    def load_docx(self, file_path: str) -> Dict:
        """
        Đọc file Word (.docx)

        Args:
            file_path: Đường dẫn file Word

        Returns:
            Dict: {'content': str, 'metadata': dict}
        """
        if not HAS_DOCX:
            raise ImportError(
                "Chua cai dat python-docx! Run: pip install python-docx")

        try:
            # Giao thức đọc file Word theo từng đoạn văn (Paragraphs)
            doc = DocxDocument(file_path)

            content = "\n".join(
                [para.text for para in doc.paragraphs if para.text.strip()])

            metadata = {
                'source': Path(file_path).name,
                'file_path': str(file_path),
                'type': 'docx',
                'paragraphs': len(doc.paragraphs),
                'format': 'Word'
            }

            print(f"Da doc Word: {Path(file_path).name}")

            return {
                'content': content.strip(),
                'metadata': metadata
            }

        except Exception as e:
            print(f"Loi doc Word {file_path}: {e}")
            return {'content': '', 'metadata': {'error': str(e)}}

    def load_text(self, file_path: str) -> Dict:
        """
        Đọc file text (.txt, .md)

        Args:
            file_path: Đường dẫn file text

        Returns:
            Dict: {'content': str, 'metadata': dict}
        """
        try:
            # Ép buộc mã hóa UTF-8 để đảm bảo đọc đúng tiếng Việt có dấu
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            metadata = {
                'source': Path(file_path).name,
                'file_path': str(file_path),
                'type': Path(file_path).suffix[1:],
                'format': 'Text'
            }

            print(f"Da doc text: {Path(file_path).name}")

            return {
                'content': content.strip(),
                'metadata': metadata
            }

        except Exception as e:
            print(f"Loi doc text {file_path}: {e}")
            return {'content': '', 'metadata': {'error': str(e)}}

    # ============================================
    # FACTORY METHOD PATTERN (MẪU THIẾT KẾ NHÀ MÁY)
    # ============================================
    # Phương thức này hoạt động như một "Bộ định tuyến" (Router).
    # Nó phân tích đuôi mở rộng của file và tự động gọi hàm xử lý tương ứng.
    def load_file(self, file_path: str) -> Dict:
        """
        Đọc file tự động dựa trên extension

        Args:
            file_path: Đường dẫn file

        Returns:
            Dict: {'content': str, 'metadata': dict}
        """
        file_path = Path(file_path)

        if not file_path.exists():
            print(f"File khong ton tai: {file_path}")
            return {'content': '', 'metadata': {'error': 'File not found'}}

        extension = file_path.suffix.lower()

        if extension == '.pdf':
            return self.load_pdf(str(file_path))
        elif extension in ['.docx', '.doc']:
            return self.load_docx(str(file_path))
        elif extension in ['.txt', '.md']:
            return self.load_text(str(file_path))
        else:
            print(f"Khong ho tro dinh dang: {extension}")
            return {'content': '', 'metadata': {'error': f'Unsupported format: {extension}'}}

    # Hàm xử lý hàng loạt (Batch Processing) cho phép đọc toàn bộ dữ liệu
    # trong một thư mục chỉ với một lần gọi hàm.
    def load_directory(
        self,
        directory: str,
        recursive: bool = True
    ) -> List[Dict]:
        """
        Đọc tất cả file trong thư mục

        Args:
            directory: Đường dẫn thư mục
            recursive: Có đọc thư mục con hay không (Duyệt đệ quy)

        Returns:
            List[Dict]: Danh sách documents
        """
        directory = Path(directory)

        if not directory.exists():
            print(f"Thu muc khong ton tai: {directory}")
            return []

        documents = []

        # Xây dựng chuỗi tìm kiếm đại diện (Glob Pattern)
        # "**/*" cho phép hệ thống đào sâu vào các thư mục con bên trong.
        pattern = "**/*" if recursive else "*"

        for file_path in directory.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
                doc = self.load_file(str(file_path))
                if doc['content']:  # Bộ lọc vệ sinh dữ liệu: Bỏ qua các file rỗng
                    documents.append(doc)

        print("\nTong ket:")
        print(f"  Da doc {len(documents)} files tu {directory}")

        return documents
